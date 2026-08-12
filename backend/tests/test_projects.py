from io import BytesIO

from docx import Document

from tests.conftest import auth_header


def _create_project(client, token: str) -> int:
    response = client.post(
        "/api/v1/projects",
        headers=auth_header(token),
        json={"name": "Test Project", "description": "d", "translation": "ESV"},
    )
    assert response.status_code == 201
    return response.json()["id"]


def test_project_crud_and_item_flow(client, token) -> None:
    headers = auth_header(token)
    project_id = _create_project(client, token)

    response = client.get("/api/v1/projects", headers=headers)
    assert len(response.json()) == 1

    item_response = client.post(
        f"/api/v1/projects/{project_id}/items",
        headers=headers,
        json={"title": "Faith and Works", "passage": "James 2:14-26", "content_type": "study_note"},
    )
    assert item_response.status_code == 201
    item_id = item_response.json()["id"]

    version_response = client.post(
        f"/api/v1/projects/{project_id}/items/{item_id}/versions",
        headers=headers,
        json={"body": "Faith without works is dead.", "change_note": "first edit"},
    )
    assert version_response.status_code == 201
    assert version_response.json()["version_number"] == 2

    versions = client.get(
        f"/api/v1/projects/{project_id}/items/{item_id}/versions", headers=headers
    )
    assert len(versions.json()) == 2

    review = client.post(
        f"/api/v1/projects/{project_id}/items/{item_id}/review",
        headers=headers,
        json={"action": "approve"},
    )
    assert review.json()["status"] == "approved"

    export = client.get(
        f"/api/v1/projects/{project_id}/items/{item_id}/export", headers=headers
    )
    assert export.status_code == 200
    assert "Faith without works is dead." in export.text

    comment = client.post(
        f"/api/v1/projects/{project_id}/items/{item_id}/comments",
        headers=headers,
        json={"body": "Check this passage context."},
    )
    assert comment.status_code == 201


def test_export_docx_format(client, token) -> None:
    headers = auth_header(token)
    project_id = _create_project(client, token)

    item_response = client.post(
        f"/api/v1/projects/{project_id}/items",
        headers=headers,
        json={"title": "Faith and Works", "passage": "James 2:14-26", "content_type": "study_note"},
    )
    item_id = item_response.json()["id"]

    client.post(
        f"/api/v1/projects/{project_id}/items/{item_id}/versions",
        headers=headers,
        json={"body": "Faith without works is dead."},
    )

    response = client.get(
        f"/api/v1/projects/{project_id}/items/{item_id}/export",
        params={"format": "docx"},
        headers=headers,
    )
    assert response.status_code == 200
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment; filename=" in response.headers["content-disposition"]
    assert response.headers["content-disposition"].endswith(".docx\"")
    assert response.content[:2] == b"PK"

    document = Document(BytesIO(response.content))
    assert document.paragraphs[0].text == "Faith and Works"
    assert any(p.text == "Faith without works is dead." for p in document.paragraphs)


def test_export_unsupported_format(client, token) -> None:
    headers = auth_header(token)
    project_id = _create_project(client, token)

    item_response = client.post(
        f"/api/v1/projects/{project_id}/items",
        headers=headers,
        json={"title": "Faith and Works", "passage": "James 2:14-26", "content_type": "study_note"},
    )
    item_id = item_response.json()["id"]

    response = client.get(
        f"/api/v1/projects/{project_id}/items/{item_id}/export",
        params={"format": "pdf"},
        headers=headers,
    )
    assert response.status_code == 400


def test_other_users_cannot_access_project(client, token) -> None:
    project_id = _create_project(client, token)

    second = {
        "email": "other@test.ai",
        "password": "other-password-1",
        "full_name": "Other User",
    }
    client.post("/api/v1/auth/register", json=second)
    login = client.post("/api/v1/auth/login", json=second)
    other_headers = auth_header(login.json()["access_token"])

    response = client.get(f"/api/v1/projects/{project_id}", headers=other_headers)
    assert response.status_code == 404
