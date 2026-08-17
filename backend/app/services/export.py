from datetime import datetime, timezone
from io import BytesIO

from docx import Document


def _render_footnotes(footnotes) -> str:
    if not footnotes:
        return ""
    lines = ["", "## Footnotes", ""]
    for i, note in enumerate(footnotes, start=1):
        if isinstance(note, dict):
            text = note.get("text", "")
        else:
            text = str(note)
        lines.append(f"{i}. {text}")
    lines.append("")
    return "\n".join(lines)


def _render_cross_refs(cross_refs) -> str:
    if not cross_refs:
        return ""
    lines = ["", "## See Also", ""]
    for ref in cross_refs:
        lines.append(f"- {ref}")
    lines.append("")
    return "\n".join(lines)


def export_markdown(
    title: str,
    body: str,
    passage: str | None = None,
    footnotes: list | None = None,
    cross_refs: list | None = None,
    status: str | None = None,
) -> str:
    parts = [f"# {title}", ""]
    meta = []
    if passage:
        meta.append(f"**Passage:** {passage}")
    if status:
        meta.append(f"**Status:** {status}")
    meta.append(f"**Exported:** {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    parts.append("\n".join(meta))
    parts.append("\n---\n")
    parts.append(body.strip())
    parts.append("\n\n---")
    parts.append(_render_footnotes(footnotes))
    parts.append(_render_cross_refs(cross_refs))
    return "\n".join(parts).strip() + "\n"


def _add_footnote_paragraphs(document: Document, footnotes) -> None:
    notes = footnotes or []
    if not notes:
        return
    document.add_paragraph("")
    document.add_heading("Footnotes", level=2)
    for i, note in enumerate(notes, start=1):
        text = note.get("text", "") if isinstance(note, dict) else str(note)
        document.add_paragraph(f"{i}. {text}")


def _add_cross_ref_paragraphs(document: Document, cross_refs) -> None:
    refs = cross_refs or []
    if not refs:
        return
    document.add_paragraph("")
    document.add_heading("See Also", level=2)
    for ref in refs:
        document.add_paragraph(f"\u2022 {ref}", style="List Bullet")


def export_docx(
    title: str,
    body: str,
    passage: str | None = None,
    footnotes: list | None = None,
    cross_refs: list | None = None,
    status: str | None = None,
) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    meta_lines = []
    if passage:
        meta_lines.append(f"Passage: {passage}")
    if status:
        meta_lines.append(f"Status: {status}")
    meta_lines.append(f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    document.add_paragraph("\n".join(meta_lines))
    document.add_paragraph("\u2500" * 40)
    for paragraph in body.splitlines():
        if paragraph.strip():
            document.add_paragraph(paragraph)
    document.add_paragraph("\u2500" * 40)
    _add_footnote_paragraphs(document, footnotes)
    _add_cross_ref_paragraphs(document, cross_refs)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
